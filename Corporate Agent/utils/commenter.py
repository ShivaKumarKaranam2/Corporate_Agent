from docx import Document
from docx.shared import RGBColor

def comment_on_docx(path: str, flags: list) -> str:
    """
    Append review notes (red colored text) to the document.
    Returns path to reviewed document.
    """
    try:
        doc = Document(path)
    except Exception:
        # if python-docx fails to open, just return original path
        return path

    if not flags:
        # nothing to add; return original path for consistency
        return path

    # Add a heading for review notes
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run("=== REVIEW NOTES ===")
    run.bold = True

    for i, flag in enumerate(flags, start=1):
        issue = flag.get("issue") or flag.get("description") or "Issue"
        suggestion = flag.get("suggestion") or ""
        para = doc.add_paragraph()
        run = para.add_run(f"{i}. {issue} — Suggestion: {suggestion}")
        # make the annotation red to stand out
        try:
            run.font.color.rgb = RGBColor(0xFF, 0x33, 0x33)
        except Exception:
            # if color fails (some environments), ignore
            pass

    reviewed_path = path.replace(".docx", "_reviewed.docx")
    doc.save(reviewed_path)
    return reviewed_path
